import os
import gradio as gr
from openai import OpenAI
from PyPDF2 import PdfReader
from docx import Document
import re

# Securely get the API key from environment variables
api_key = os.getenv("OPENAI_API_KEY")
base_url = "https://api.aimlapi.com/v1"
llama_model = "meta-llama/Meta-Llama-3.1-70B-Instruct-Turbo"

# Initialize the OpenAI API
api = OpenAI(api_key=api_key, base_url=base_url)

# Function to extract text from a PDF file
def extract_text_from_pdf(pdf_file):
    reader = PdfReader(pdf_file)
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    return text

# Function to extract user details from CV (e.g., Name, Contact Info, etc.)
def extract_user_details(cv_text):
    lines = cv_text.splitlines()
    name = lines[0] if len(lines) > 0 else "Name Not Found"
    title = lines[1] if len(lines) > 1 else "Title Not Found"
    contact_info = lines[2] if len(lines) > 2 else "Contact Info Not Found"
    
    return name, title, contact_info

# Function to match and regenerate CV content
def regenerate_cv(cv_text, job_description):
    system_prompt = "You are a world class pro CV writing assistant."
    user_prompt = (f"Here is a CV: {cv_text}\n\n"
                   f"And here is the job description: {job_description}\n\n"
                   "Please regenerate the CV according to the job description, "
                   "fill the gaps, add all sufficient skills that match the job description and provide it in a professional US template that has the highest possible ATS score.")

    completion = api.chat.completions.create(
        model=llama_model,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.7,
        max_tokens=4096,  # Increase max_tokens to handle larger outputs
    )

    response = completion.choices[0].message.content
    return response

# Function to generate the interview preparation note
def generate_interview_note(cv_text, updated_cv_text):
    system_prompt = "You are a world class pro interview preparation assistant."
    user_prompt = (f"Here is the original CV: {cv_text}\n\n"
                   f"Here is the updated CV: {updated_cv_text}\n\n"
                   "Identify the changes made and provide a list of topics the candidate should prepare for the interview based on the updates in a very concise and professional manner with proper headings and bullet points.")

    completion = api.chat.completions.create(
        model=llama_model,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.7,
        max_tokens=1024,
    )

    response = completion.choices[0].message.content
    return response

# Function to calculate ATS score
def calculate_ats_score(cv_text, job_description):
    job_keywords = re.findall(r'\b\w+\b', job_description.lower())
    cv_keywords = re.findall(r'\b\w+\b', cv_text.lower())
    
    matching_keywords = set(job_keywords) & set(cv_keywords)
    ats_score = len(matching_keywords) / len(set(job_keywords)) * 100
    return round(ats_score, 2)

# Function to create and save the regenerated CV as DOCX (formatted)
def create_formatted_cv(cv_text, file_name="Updated_CV.docx"):
    doc = Document()

    # Extract user details (name, title, contact info)
    name, title, contact_info = extract_user_details(cv_text)

    # Adding CV content with formatting
    doc.add_heading(name, level=1)
    doc.add_paragraph(title)
    doc.add_paragraph(contact_info)

    # Split the CV into sections
    sections = cv_text.split('\n\n')

    # Define headings for each section based on the template
    headings = ["ACHIEVEMENTS", "EDUCATION", "SKILLS", "WORK EXPERIENCE", "VOLUNTEER EXPERIENCE", "HACKATHON PROJECTS", "WORKSHOPS AND WEBINARS"]

    for section in sections:
        for heading in headings:
            if heading in section.upper():
                doc.add_heading(heading, level=2)
        doc.add_paragraph(section)

    # Save the document
    doc.save(file_name)
    return file_name

# Function to process the CV and generate outputs
def process_cv(cv_pdf, job_description):
    # Extract text from the uploaded CV PDF
    cv_text = extract_text_from_pdf(cv_pdf)

    # Calculate initial ATS score
    initial_ats_score = calculate_ats_score(cv_text, job_description)

    # Regenerate the CV
    updated_cv_text = regenerate_cv(cv_text, job_description)

    # Calculate updated ATS score
    updated_ats_score = calculate_ats_score(updated_cv_text, job_description)

    # Generate interview preparation note
    interview_note = generate_interview_note(cv_text, updated_cv_text)

    # Create and save the formatted CV
    formatted_cv_path = create_formatted_cv(updated_cv_text)

    # Create the interview notes file
    interview_notes_path = "Interview_Notes.txt"
    with open(interview_notes_path, "w") as f:
        f.write(interview_note)
    
    return updated_cv_text, initial_ats_score, updated_ats_score, interview_note, formatted_cv_path, interview_notes_path

# Define the Gradio app
def app_interface():
    with gr.Blocks() as interface:
        gr.Markdown("### **HireFit** By team Mixed Intelligence")

        # Short description
        gr.Markdown("""
        **Version 1.0**
        This project takes your CV and job description, then provides a new CV optimized for the specific job description. It also highlights gaps in your CV and provides detailed interview preparation notes to help you succeed.
        """)

        with gr.Row():
            # Left panel
            with gr.Column(scale=1):
                gr.Markdown("### Upload and Generate")
                cv_pdf = gr.File(label="Upload Your CV (PDF)")
                job_description = gr.Textbox(label="Paste Job Description")
                process_btn = gr.Button("Regenerate CV", elem_id="process_btn")

            # Middle panel
            with gr.Column(scale=2):
                gr.Markdown("### Regenerated CV and Interview Notes")
                updated_cv_text = gr.Textbox(label="Regenerated CV", lines=20, interactive=False)
                interview_note = gr.Textbox(label="Interview Preparation Notes", lines=10, interactive=False)

            # Right panel
            with gr.Column(scale=1):
                gr.Markdown("### ATS Scores")
                ats_score_before = gr.Number(label="Original ATS Score", value=0)
                ats_score_after = gr.Number(label="Updated ATS Score", value=0)

                # Adding download buttons to the right panel
                download_cv_btn = gr.File(label="Download Updated CV", elem_id="download_cv_btn")
                download_notes_btn = gr.File(label="Download Interview Notes", elem_id="download_notes_btn")

        # Button click event
        process_btn.click(
            fn=process_cv,
            inputs=[cv_pdf, job_description],
            outputs=[updated_cv_text, ats_score_before, ats_score_after, interview_note, download_cv_btn, download_notes_btn]
        )

    return interface

if __name__ == "__main__":
    app_interface().launch()
